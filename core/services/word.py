import re
from copy import deepcopy
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from docx.shared import Pt, RGBColor
from fastapi import UploadFile

from core.services.base import BaseDocumentService
import docx
import markdown
import html
from bs4 import BeautifulSoup
from docx import Document


class WordService(BaseDocumentService):
    def __init__(self):
        self.docx_file = None

    def load(self, file) -> None:
        self.docx_file = Document(file)


    def update(self, params: dict) -> None:
        if not self.docx_file:
            raise ValueError("Word файл не загружен.")
        
        result_doc = Document()

        #проход по всем элементам(включая таблицы и тд)
        for index, block in enumerate(params['blocks']):
            doc, temp_filename = self.copy_to_temp(index)
            for key, value in block.items():
                bookmark_found = False
                if isinstance(value, str):
                    for element in doc.element.body.iter():
                        if element.tag == qn('w:bookmarkStart'):  # тег закладок для поиска в списке xml
                            bookmark_name = element.get(qn('w:name'))
                            if bookmark_name == key:
                                self.replace_text(doc, element, value)
                                bookmark_found = True
                    if not bookmark_found:
                        print(f"Закладка {key} в документе не найдена")
                elif isinstance(value, list):
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                bookmark_element = self.find_bookmark_in_table(cell, key)
                                if bookmark_element != None:
                                    print(f"мы сейчас на ключе {key}")
                                    for paragraph in cell.paragraphs:
                                        if paragraph.runs:
                                            for run in paragraph.runs:
                                                font = run.font
                                                run.text = value[0]
                                                print(f"вставка в {key} произошла успешно")
                                                run.font.bold = font.bold
                                                run.font.italic = font.italic
                                                run.font.underline = font.underline
                                                run.font.color.rgb = font.color.rgb
                                                run.font.name = font.name
                                                run.font.size = font.size
                                                bookmark_found = True
                                                start_index = None
                                            
                                                for idx, c in enumerate(row.cells):
                                                    if c.text.strip() == cell.text.strip():
                                                        start_index = idx
                                                        break

                                                if start_index is not None:
                                                    for i, val in enumerate(value[1:], start=1):
                                                        if start_index + i < len(row.cells):
                                                            row.cells[start_index + i].text = val
                                                break
                                        else:
                                            run = paragraph.add_run()
                                            font = run.font
                                            run.text = value[0]
                                            print(f"вставка в {key} произошла успешно")
                                            run.font.bold = font.bold
                                            run.font.italic = font.italic
                                            run.font.underline = font.underline
                                            run.font.color.rgb = font.color.rgb
                                            run.font.name = font.name
                                            run.font.size = font.size
                                            bookmark_found = True
                                            start_index = None

                                            for idx, c in enumerate(row.cells):
                                                if c.text.strip() == cell.text.strip():
                                                    start_index = idx
                                                    break

                                            if start_index is not None:
                                                for i, val in enumerate(value[1:], start=1):
                                                    if start_index + i < len(row.cells):
                                                        row.cells[start_index + i].text = val
                                            break


                            if bookmark_found:
                                break
                        if bookmark_found:
                            break

            doc.save(temp_filename)
            self.add_temp_to_original(result_doc, temp_filename, params, index)

            #удаление временных файлов
            if os.path.exists(temp_filename):
                os.remove(temp_filename)

        if params['newpage'] == 'true':
            last_paragraph = result_doc.paragraphs[-1]
            p = last_paragraph._element
            p.getparent().remove(p)

        self.docx_file = result_doc

    def find_bookmark_in_table(self, cell, bookmark_name):
        for paragraph in cell.paragraphs:
            for element in paragraph._element.iter():
                if element.tag == qn('w:bookmarkStart'):
                    if element.get(qn('w:name')) == bookmark_name:
                        return element  #возвращаем элемент закладки
        return None  #если закладка не найдена

    def copy_to_temp(self, index):
        #функция копирования данных во временные файлы
        temp_filename = f'temp{index}.docx'
        self.docx_file.save(temp_filename)
        new_doc = Document(temp_filename)
        return new_doc, temp_filename

    def add_temp_to_original(self, original_doc, temp_doc_path, params: dict, index):
        # Открытие временного документа
        temp_doc = Document(temp_doc_path)

        elements_to_copy = list(temp_doc.element.body)
        paragraph_index = 0
        table_index = 0

        for element in elements_to_copy:
            # Если элемент — это параграф
            if element.tag.endswith('p'):
                if paragraph_index < len(temp_doc.paragraphs):
                    paragraph = temp_doc.paragraphs[paragraph_index]
                    new_paragraph = original_doc.add_paragraph()

                    # Копирование форматирования абзаца
                    self.copy_paragraph_formatting(paragraph, new_paragraph)

                    # Копирование всех runs внутри абзаца
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        self.copy_run_formatting(run, new_run)

                    paragraph_index += 1

            # Если элемент — это таблица
            elif element.tag.endswith('tbl'):
                if table_index < len(temp_doc.tables):
                    table = temp_doc.tables[table_index]

                    # Копирование таблицы в оригинальный документ
                    new_table = original_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style

                    # Копирование ячеек и их содержимого
                    for row_index, row in enumerate(table.rows):
                        for col_index, cell in enumerate(row.cells):
                            new_cell = new_table.cell(row_index, col_index)
                            new_cell.vertical_alignment = cell.vertical_alignment
                            for paragraph in cell.paragraphs:
                                new_paragraph = new_cell.add_paragraph()
                                self.copy_paragraph_formatting(paragraph,
                                                               new_paragraph)  # Копирование форматирования абзаца
                                for run in paragraph.runs:
                                    new_run = new_paragraph.add_run(run.text)
                                    self.copy_run_formatting(run, new_run)  # Копирование форматирования внутри "run"

                            # Дополнительное форматирование ячеек (например, границы)
                            tcPr = new_cell._tc.get_or_add_tcPr()
                            tcBorders = OxmlElement("w:tcBorders")
                            for border in ["top", "left", "bottom", "right"]:
                                element = OxmlElement(f"w:{border}")
                                element.set(qn("w:val"), "single")
                                element.set(qn("w:sz"), "4")
                                element.set(qn("w:space"), "0")
                                element.set(qn("w:color"), "auto")
                                tcBorders.append(element)
                            tcPr.append(tcBorders)

                    table_index += 1

        # Если требуется новая страница, добавляем разрыв страницы
        if params.get('newpage', 'false') == 'true':
            original_doc.add_page_break()

    def copy_paragraph_formatting(self, paragraph, new_paragraph):
        # Копирование форматирования абзаца
        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
        new_paragraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing

    def copy_run_formatting(self, run, new_run):
        # Копирование форматирования run (жирный, курсив, цвет и т.д.)
        new_run.font.bold = run.font.bold
        new_run.font.italic = run.font.italic
        new_run.font.underline = run.font.underline
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.strike = run.font.strike
        new_run.font.highlight_color = run.font.highlight_color


    def replace_text(self, doc, bookmark_element, new_text):
        parent_element = bookmark_element.getparent()
        for sibling in bookmark_element.itersiblings():
            if sibling.tag == qn('w:r'):
                text_elements = sibling.findall(qn('w:t'))
                if text_elements:
                    new_run = deepcopy(sibling)
                    for text_elem in new_run.findall(qn('w:t')):
                        text_elem.text = new_text

                    parent_element.insert(parent_element.index(sibling), new_run)
                    parent_element.remove(sibling)
                    parent_element.remove(bookmark_element)
                    return

        new_run = OxmlElement('w:r')
        new_text_element = OxmlElement('w:t')
        new_text_element.text = new_text
        new_run.append(new_text_element)
        parent_element.insert(parent_element.index(bookmark_element), new_run)
        parent_element.remove(bookmark_element)



    def copy_paragraph(self, paragraph, document):
        #ункция для копирования абзаца в новый документ
        new_paragraph = document.add_paragraph()
        #new_paragraph.style = paragraph.style  #стиль абзаца
        new_paragraph.style.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)  #текст
            new_run.bold = run.bold  #жирность
            new_run.italic = run.italic  #курсив
            new_run.font.size = run.font.size  #размер шрифта
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb  #цвет шрифта(если есть)

    def save_to_bytes(self) -> BytesIO:
        if not self.docx_file:
            raise ValueError("Word файл не загружен.")
        output = BytesIO()
        self.docx_file.save(output)
        output.seek(0)
        return output

    def save_to_file(self, file_path: str) -> None:
        if self.docx_file:
            self.docx_file.save(file_path)
        else:
            raise ValueError("Word файл не загружен.")

    def clean_para_with_bookmark(self, params: dict) -> None:
        if not self.docx_file:
            raise ValueError("Word файл не загружен.")

        bookmark_id = 0
        if 'bookmarks' in params and isinstance(params['bookmarks'], list):
            for bookmark in params['bookmarks']:
                if isinstance(bookmark, dict):
                    for text_to_remove, key in bookmark.items():
                        self.process_doc(text_to_remove, key, self.docx_file, bookmark_id)

    def process_doc(self, text_to_remove, key, doc, bookmark_id):
        for para in doc.paragraphs:
            if text_to_remove in para.text:
                para.text = para.text.replace(text_to_remove, '')
                self.add_bookmark(para, key, bookmark_id)
                bookmark_id += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text_to_remove in para.text:
                            para.text = para.text.replace(text_to_remove, '')
                            self.add_bookmark(para, key, bookmark_id)
                            bookmark_id += 1

    def add_bookmark(self, para, key, bookmark_id):
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(bookmark_id))
        bookmark_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name', key)

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(bookmark_id))

        run = para.add_run()
        run._element.append(bookmark_start)
        run._element.append(bookmark_end)

    def extract_bookmarks(self) -> dict:
        template = {}
        #проход по всем элементам(включая таблицы и тд)
        for element in self.docx_file.element.body.iter():
            if element.tag == qn('w:bookmarkStart'): #тег закладок для поиска в xml
                bookmark_name = element.get(qn('w:name'))
                if bookmark_name and bookmark_name != "_GoBack":
                    text = self.get_text_from_bookmarks(self.docx_file, element)
                    template[bookmark_name] = text

        output_data = {
            "blocks": [template],
            "newpage": "false"  # Список всех шаблонов с массивами
        }
        return output_data

    #получаем текст с закладок
    def get_text_from_bookmarks(self, doc, bookmark_start):
        bookmark_text = []
        inside_bookmark = False
        start_id = bookmark_start.get(qn('w:id'))

        for element in doc.element.body.iter():
            #начало закладки
            if element.tag == qn('w:bookmarkStart') and element.get(qn('w:id')) == start_id:
                inside_bookmark = True

            #собираем текст между началом и концом закладки
            if inside_bookmark and element.tag == qn('w:t'):
                bookmark_text.append(element.text)

            #конец закладки
            if element.tag == qn('w:bookmarkEnd') and element.get(qn('w:id')) == start_id:
                #брейкаем после конца закладки
                break

        return ' '.join(bookmark_text).strip()


    def parse_with_formatting(self) -> dict:
        blocks = []
        block_template = {}

        for element in self.docx_file.element.body.iter():
            if element.tag == qn('w:bookmarkStart'):
                bookmark_name = element.get(qn('w:name'))
                if bookmark_name and bookmark_name != "_GoBack":
                    text, format_info = self.get_text_and_format_from_bookmark(self.docx_file, element)
                    if isinstance(text, list):
                        block_template[bookmark_name] = text
                    else:
                        if text == "":
                            block_template[bookmark_name] = {
                                "value": "",
                                "format": None
                            }
                        else:
                            block_template[bookmark_name] = {
                                "value": text,
                                "format": format_info
                            }

        blocks.append(block_template)
        output_data = {
            "blocks": blocks,
            "newpage": "false"
        }
        return output_data

    def get_text_and_format_from_bookmark(self, doc, bookmark_element):
        text_content = []
        format_info = None
        inside_bookmark = False
        start_id = bookmark_element.get(qn('w:id'))

        for element in doc.element.body.iter():
            if element.tag == qn('w:bookmarkStart') and element.get(qn('w:id')) == start_id:
                inside_bookmark = True

            if inside_bookmark:
                if element.tag == qn('w:t'):
                    text_content.append(element.text if element.text else "")

                if element.tag == qn('w:bookmarkEnd') and element.get(qn('w:id')) == start_id:
                    inside_bookmark = False
                    break

            if inside_bookmark and element.tag.endswith('r'):
                run_format = self.extract_run_format(element)
                if run_format:
                    format_info = run_format

        text = ' '.join(text_content).strip()
        if len(text_content) > 1:
            return text_content, format_info
        return text, format_info

    def extract_run_format(self, run):
        font = run.find(".//w:rPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        if font is None or run.text is None:
            return None

        format_info = {
            "fontname": "Calibri",
            "fontsize": "11",
            "fillcolor": None,
            "textcolor": None,
            "bold": False,
            "italic": False,
            "underline": False,
            "strikethrough": False,
            "align": "left",
        }

        if font.find(qn('w:b')) is not None:
            format_info["bold"] = True

        if font.find(qn('w:i')) is not None:
            format_info["italic"] = True

        if font.find(qn('w:u')) is not None:
            format_info["underline"] = True

        if font.find(qn('w:strike')) is not None:
            format_info["strikethrough"] = True

        font_size = font.find(qn('w:sz'))
        if font_size is not None:
            format_info["fontsize"] = int(font_size.get(qn('w:val'))) / 2

        font_color = font.find(qn('w:color'))
        if font_color is not None:
            format_info["textcolor"] = font_color.get(qn('w:val'))

        font_fill = font.find(qn('w:highlight'))
        if font_fill is not None:
            format_info["fillcolor"] = font_fill.get(qn('w:val'))

        font_name = font.find(qn('w:rFonts'))
        if font_name is not None:
            format_info["fontname"] = font_name.get(qn('w:ascii'))

        parent_paragraph = run.getparent().getparent()
        if parent_paragraph is not None:
            alignment = parent_paragraph.find(".//w:jc", namespaces={
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if alignment is not None:
                align_val = alignment.get(qn('w:val'))
                if align_val in ['left', 'center', 'right', 'both']:
                    format_info["align"] = align_val
        return format_info


    def update_with_formatting(self, params: dict) -> None:
        if not self.docx_file:
            raise ValueError("Word файл не загружен.")
        result_doc = Document()
        for index, block in enumerate(params['blocks']):
            doc, temp_filename = self.copy_to_temp(index)
            for key, value in block.items():
                value_for_check = {}
                if isinstance(value, dict):
                    value_for_check = value.get('value')
                bookmark_found = False
                if isinstance(value, dict) and isinstance(value_for_check, str) and 'value' in value and 'format' in value:
                    print(f"{key} мы в dict")
                    new_text = value['value']
                    formatting = value['format']
                    for element in doc.element.body.iter():
                        if element.tag == qn('w:bookmarkStart'):
                            bookmark_name = element.get(qn('w:name'))
                            if bookmark_name == key:
                                self.replace_text_for_formatting(doc, element, new_text, formatting)
                                bookmark_found = True
                    if not bookmark_found:
                        print(f"Закладка {key} в документе не найдена")
                elif isinstance(value, str):
                    print(f"{key} мы в str")
                    for element in doc.element.body.iter():
                        if element.tag == qn('w:bookmarkStart'):
                            bookmark_name = element.get(qn('w:name'))
                            if bookmark_name == key:
                                self.replace_text(doc, element, value)
                                bookmark_found = True
                    if not bookmark_found:
                        print(f"Закладка {key} в документе не найдена")

                elif isinstance(value, dict) and isinstance(value_for_check, list):
                    print(f"{key} мы в list")
                    formatting = None
                    if isinstance(value, dict):
                        formatting = value.get('format')
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                bookmark_element = self.find_bookmark_in_table(cell, key)
                                if bookmark_element is not None:
                                    for paragraph in cell.paragraphs:
                                        if paragraph.runs:
                                            for run in paragraph.runs:
                                                run.text = value_for_check[0]
                                                self.apply_formatting(run, formatting, paragraph)
                                                bookmark_found = True
                                                start_index = None

                                                for idx, c in enumerate(row.cells):
                                                    if c.text == cell.text:
                                                        start_index = idx
                                                        break

                                                if start_index is not None:
                                                    for i, val in enumerate(value_for_check[1:], start=1):
                                                        if start_index + i < len(row.cells):
                                                            row.cells[start_index + i].text = val
                                                            new_cell = row.cells[start_index + i]
                                                            for paragraph in new_cell.paragraphs:
                                                                for run in paragraph.runs:
                                                                    self.apply_formatting(run, formatting, paragraph)
                                                    break
                                        else:
                                            run = paragraph.add_run()
                                            run.text = value_for_check[0]
                                            self.apply_formatting(run, formatting, paragraph)
                                            bookmark_found = True
                                            start_index = None

                                            for idx, c in enumerate(row.cells):
                                                if c.text == cell.text:
                                                    start_index = idx
                                                    break

                                            if start_index is not None:
                                                for i, val in enumerate(value_for_check[1:], start=1):
                                                    if start_index + i < len(row.cells):
                                                        row.cells[start_index + i].text = val
                                                        new_cell = row.cells[start_index + i]
                                                        for paragraph in new_cell.paragraphs:
                                                            for run in paragraph.runs:
                                                                self.apply_formatting(run, formatting, paragraph)
                                                break

                                if bookmark_found:
                                    break
                        if bookmark_found:
                            break


                elif isinstance(value, list):
                    print(f"{key} мы в list без форматирования")
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                bookmark_element = self.find_bookmark_in_table(cell, key)
                                if bookmark_element != None:
                                    for paragraph in cell.paragraphs:
                                        if paragraph.runs:
                                            for run in paragraph.runs:
                                                font = run.font
                                                run.text = value[0]
                                                run.font.bold = font.bold
                                                run.font.italic = font.italic
                                                run.font.underline = font.underline
                                                run.font.color.rgb = font.color.rgb
                                                run.font.name = font.name
                                                run.font.size = font.size
                                                bookmark_found = True
                                                start_index = None
                                                for idx, c in enumerate(row.cells):
                                                    if c.text.strip() == cell.text.strip():
                                                        start_index = idx
                                                        break
                                                if start_index is not None:
                                                    for i, val in enumerate(value[1:], start=1):
                                                        if start_index + i < len(row.cells):
                                                            row.cells[start_index + i].text = val
                                                break
                                        else:
                                            run = paragraph.add_run()
                                            font = run.font
                                            run.text = value[0]
                                            run.font.bold = font.bold
                                            run.font.italic = font.italic
                                            run.font.underline = font.underline
                                            run.font.color.rgb = font.color.rgb
                                            run.font.name = font.name
                                            run.font.size = font.size
                                            bookmark_found = True
                                            start_index = None

                                            for idx, c in enumerate(row.cells):
                                                if c.text.strip() == cell.text.strip():
                                                    start_index = idx
                                                    break

                                            if start_index is not None:
                                                for i, val in enumerate(value[1:], start=1):
                                                    if start_index + i < len(row.cells):
                                                        row.cells[start_index + i].text = val
                                            break

                                if bookmark_found:
                                    break
                        if bookmark_found:
                            break

                    if not bookmark_found:
                        print(f"Закладка '{key}' не найдена в документе.")

            doc.save(temp_filename)
            self.add_temp_to_original(result_doc, temp_filename, params, index)

            if os.path.exists(temp_filename):
                os.remove(temp_filename)

        if params['newpage'] == 'true':
            last_paragraph = result_doc.paragraphs[-1]
            p = last_paragraph._element
            p.getparent().remove(p)

        self.docx_file = result_doc

    def replace_text_for_formatting(self, doc, bookmark_element, new_text, formatting=None):
        parent_element = bookmark_element.getparent()

        for sibling in bookmark_element.itersiblings():
            if sibling.tag == qn('w:r'):
                text_elements = sibling.findall(qn('w:t'))
                if text_elements:
                    paragraph = doc.add_paragraph()
                    new_run = paragraph.add_run()

                    old_text = ''.join([t.text for t in text_elements if t.text])
                    new_run.text = new_text if new_text else old_text

                    self.apply_formatting(new_run, formatting, paragraph)

                    parent_element.insert(
                        parent_element.index(sibling),
                        new_run._element
                    )

                    parent_element.remove(sibling)
                    parent_element.remove(bookmark_element)
                    return

        paragraph = doc.add_paragraph()
        new_run = paragraph.add_run(new_text)

        self.apply_formatting(new_run, formatting, paragraph)

        parent_element.insert(
            parent_element.index(bookmark_element),
            new_run._element
        )
        parent_element.remove(bookmark_element)

    def apply_formatting(self, run, formatting, paragraph):
        if formatting:
            if 'fontname' in formatting:
                run.font.name = formatting['fontname']

            if 'fontsize' in formatting:
                try:
                    fontsize = float(formatting['fontsize'])
                    if fontsize > 0:
                        run.font.size = Pt(fontsize)
                except (ValueError, TypeError):
                    print(f"Ошибка: 'fontsize' должен быть числом, получено {formatting['fontsize']}.")

            if 'fillcolor' in formatting:
                highlight_color = formatting['fillcolor']
                if isinstance(highlight_color, str):
                    try:
                        highlight_color = WD_COLOR_INDEX[highlight_color.upper()]
                        highlight_color_value = highlight_color.value
                        run.font.highlight_color = highlight_color_value
                    except KeyError:
                        raise ValueError(f"'{highlight_color}' не поддерживается в WD_COLOR_INDEX")
                else:
                    run.font.highlight_color = highlight_color

            if 'textcolor' in formatting:
                try:
                    rgb = self.hex_to_rgb(formatting['textcolor'])
                    run.font.color.rgb = RGBColor(*rgb)
                except ValueError as e:
                    raise ValueError(f"Ошибка при обработке цвета текста: {e}")

            if 'bold' in formatting:
                run.bold = formatting['bold']

            if 'italic' in formatting:
                run.italic = formatting['italic']

            if 'underline' in formatting:
                run.underline = formatting['underline']

            if 'strikethrough' in formatting:
                run.font.strike = formatting['strikethrough']

            if 'align' in formatting:
                if formatting['align'] == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif formatting['align'] == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif formatting['align'] == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif formatting['align'] == 'both':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def hex_to_rgb(self, hex_color):
        if hex_color is None:
            print("Ошибка: hex_color равен None.")
            return (0, 0, 0)
        hex_color = hex_color.lstrip('#')
        if len(hex_color) != 6:
            raise ValueError(f"Некорректный формат цвета: {hex_color}. Ожидается строка из 6 символов.")
        try:
            return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
        except ValueError:
            raise ValueError(f"Некорректный формат hex цвета: {hex_color}")

    def md_to_word(self, md_file: bytes) -> None:
        md_content = md_file.decode("utf-8")
        lines = md_content.splitlines()

        doc = Document()

        buffer = []
        i = 0
        while i < len(lines):
            line = lines[i]

            if '|' in line and i + 1 < len(lines) and set(lines[i + 1].strip()) <= {'-', '|', ' '}:
                if buffer:
                    self._process_md_block(doc, '\n'.join(buffer))
                    buffer = []

                table_lines = [line, lines[i + 1]]
                i += 2
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1

                headers, rows = self.parse_md_table(table_lines)
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Table Grid'

                for col_idx, header in enumerate(headers):
                    table.cell(0, col_idx).text = header

                for row_idx, row in enumerate(rows, start=1):
                    for col_idx, cell in enumerate(row):
                        table.cell(row_idx, col_idx).text = cell

                doc.add_paragraph('')
            else:
                buffer.append(line)
                i += 1

        if buffer:
            self._process_md_block(doc, '\n'.join(buffer))

        self.docx_file = doc

    def _process_md_block(self, doc, md_text: str):
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, "html.parser")
        for element in soup.contents:
            self.process_element(doc, element)

    def remove_md_tables(self, lines):
        result = []
        skip = False
        for line in lines:
            if '|' in line and not skip:
                skip = True
                continue
            if skip and ('|' in line or set(line.strip()) <= {'-', '|', ' '}):
                continue
            skip = False
            result.append(line)
        return result

    def extract_md_tables(self, md_lines):
        tables = []
        i = 0
        while i < len(md_lines):
            if '|' in md_lines[i]:
                table_lines = []
                while i < len(md_lines) and '|' in md_lines[i]:
                    table_lines.append(md_lines[i])
                    i += 1
                if len(table_lines) >= 2 and re.match(r'^\s*\|?[\s:-]+\|', table_lines[1]):
                    tables.append(table_lines)
            else:
                i += 1
        return tables

    def parse_md_table(self, table_lines):
        headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
        rows = []
        for line in table_lines[2:]:
            cols = [c.strip() for c in line.split('|') if c.strip()]
            if cols:
                rows.append(cols)
        return headers, rows


    def process_element(self, doc, element, list_style=None):
        html_to_docx_styles = {
            "h1": "Heading 1",
            "h2": "Heading 2",
            "h3": "Heading 3",
            "h4": "Heading 4",
            "blockquote": "Quote",
        }
        if element.name in html_to_docx_styles:
            self.add_formatted_paragraph(doc, element, style=html_to_docx_styles[element.name])
        elif element.name == "p":
            self.add_formatted_paragraph(doc, element)
            if element.find("img"):
                self.add_image(doc, element.find("img"))
        elif element.name == "ul":
            self.process_list(doc, element, "List Bullet", is_ordered=False)
        elif element.name == "ol":
            self.process_list(doc, element, "List Number", is_ordered=True)
        elif element.name == "li":
            self.add_list_item(doc, element, list_style)
        elif element.name == "blockquote":
            self.add_quote(doc, element)
        elif element.name == "img":
            self.add_image(doc, element)

    def process_list(self, doc, element, list_style, is_ordered):
        if is_ordered:
            for li in element.find_all("li"):
                self.process_element(doc, li, list_style="List Number")
        else:
            for li in element.find_all("li"):
                self.process_element(doc, li, list_style="List Bullet")

    def add_list_item(self, doc, element, list_style):
        paragraph = doc.add_paragraph(style=list_style)

        if element.find("p"):
            for p in element.find_all("p"):
                self.process_inline_formatting(p, paragraph)
        else:
            self.process_inline_formatting(element, paragraph)

    def add_formatted_paragraph(self, doc, element, style=None):
        paragraph = doc.add_paragraph(style=style)
        self.process_inline_formatting(element, paragraph)

    def process_inline_formatting(self, element, paragraph):
        for content in element.contents:
            if isinstance(content, str):
                paragraph.add_run(html.unescape(content))
            elif content.name == "strong":
                run = paragraph.add_run(html.unescape(content.get_text()))
                run.bold = True
            elif content.name == "em":
                run = paragraph.add_run(html.unescape(content.get_text()))
                run.italic = True
            elif content.name == "a":
                self.add_hyperlink(paragraph, content.get_text(), content["href"])
            else:
                self.process_inline_formatting(content, paragraph)

    def add_hyperlink(self, paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement('w:r'), paragraph)
        new_run.text = text
        new_run.style = self.get_or_create_hyperlink_style(part.document)

        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)

    def get_or_create_hyperlink_style(self, d):
        if "Hyperlink" not in d.styles:
            hs = d.styles.add_style("Hyperlink", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True)
            hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
            hs.font.underline = True
        return "Hyperlink"

    def add_quote(self, doc, element):
        paragraph = doc.add_paragraph(style="Quote")
        self.process_inline_formatting(element, paragraph)

    def add_image(self, doc, img_element):
        image_name = img_element.get("alt", "Изображение")
        image_url = img_element.get("src")
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{image_name}: {image_url}")
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
